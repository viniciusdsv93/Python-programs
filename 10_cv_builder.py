from docx import Document
from docx.shared import Inches

document = Document()

# welcome message
print('=' * 50)
print('WELCOME TO THE CV BUILDER'.center(50))
print('=' * 50, '\n')

# adding the profile picture
document.add_picture(
    'profile.jpeg',
    width = Inches(1.0)
)

# gathering user's data to insert on the cv
name = input('Enter your name: ')
phone_number = input('\nEnter your phone number: ')
email = input('\nEnter your email: ')

document.add_paragraph(f'{name} | {phone_number} | {email}')

# about me
document.add_heading('About me')
about_me = input('\nWrite about yourself: ')
document.add_paragraph(f'{about_me}')

# work experience
document.add_heading('\nWork experience: ')

while True:
    p = document.add_paragraph()

    company = input('\nEnter company: ')
    start_date = input('From date: ')
    end_date = input('To date: ')

    p.add_run(f'{company} ').bold = True
    p.add_run(f'{start_date} - {end_date} \n').italic = True

    details = input(f'\nDescribe your experience at {company} ')

    p.add_run(details)

    option = input('\nInsert one more? [Y/N] ').strip().upper()[0]

    if option == 'N':
        break

# skills
document.add_heading('\nSkills')

while True:
    q = document.add_paragraph()
    q.add_run(input('\nInsert your skills: '))
    q.style = 'List Bullet'
    option = input('\nDo you have any other skills to insert? [Y/N] ').strip().upper()[0]
    if option == 'N':
        break

# footer
section = document.sections[0]
footer = section.footer
r = footer.paragraphs[0]
r.text = 'CV generated using Python.'

document.save('cv.docx')
